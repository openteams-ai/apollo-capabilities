"""Document Summarizer — upload a document and summarize it via an LLM.

Supports plain text, Markdown, PDF, and Word (.docx) inputs. Talks to
three local OpenAI-compatible servers (llama.cpp, Ollama, Docker Model
Runner) or a generic remote OpenAI-compatible API, and streams the
summary back into the page.
"""
from __future__ import annotations

import io
from dataclasses import dataclass
from pathlib import Path
from typing import Callable

import requests
import streamlit as st
from docx import Document as DocxDocument
from openai import OpenAI
from pypdf import PdfReader

# ---------------------------------------------------------------------------
# Provider registry
#
# Local providers expose an OpenAI-compatible chat-completions endpoint and
# don't need an API key. The "remote" entry is a generic OpenAI-compatible
# provider — point it at any host that speaks /v1/chat/completions.
# ---------------------------------------------------------------------------


@dataclass(frozen=True)
class Provider:
    key: str
    label: str
    default_base_url: str
    chat_base_url_suffix: str
    probe: Callable[..., list[str]]
    needs_api_key: bool
    notes: str


def _probe_openai_models(api_url: str, api_key: str = "") -> list[str]:
    headers = {"Authorization": f"Bearer {api_key}"} if api_key else {}
    r = requests.get(f"{api_url.rstrip('/')}/models", timeout=5, headers=headers)
    r.raise_for_status()
    payload = r.json()
    items = payload.get("data") or payload.get("models") or []
    out: list[str] = []
    for it in items:
        if isinstance(it, str):
            out.append(it)
        elif isinstance(it, dict):
            mid = it.get("id") or it.get("name")
            if mid:
                out.append(mid)
    return sorted(out)


def probe_llamacpp(base_url: str, api_key: str = "") -> list[str]:
    return _probe_openai_models(f"{base_url.rstrip('/')}/v1")


def probe_ollama(base_url: str, api_key: str = "") -> list[str]:
    r = requests.get(f"{base_url.rstrip('/')}/api/tags", timeout=2)
    r.raise_for_status()
    return sorted(m["name"] for m in r.json().get("models", []) if m.get("name"))


def probe_docker(base_url: str, api_key: str = "") -> list[str]:
    return _probe_openai_models(f"{base_url.rstrip('/')}/engines/v1")


def probe_remote(base_url: str, api_key: str = "") -> list[str]:
    return _probe_openai_models(base_url, api_key)


PROVIDERS: dict[str, Provider] = {
    "llamacpp": Provider(
        key="llamacpp",
        label="llama.cpp (local)",
        default_base_url="http://localhost:8080",
        chat_base_url_suffix="/v1",
        probe=probe_llamacpp,
        needs_api_key=False,
        notes="Default `pixi run serve` port for the llamacpp capability is 8080.",
    ),
    "ollama": Provider(
        key="ollama",
        label="Ollama (local)",
        default_base_url="http://localhost:11434",
        chat_base_url_suffix="/v1",
        probe=probe_ollama,
        needs_api_key=False,
        notes="Start with `ollama serve`; the desktop app starts it automatically.",
    ),
    "docker": Provider(
        key="docker",
        label="Docker Model Runner (local)",
        default_base_url="http://localhost:12434",
        chat_base_url_suffix="/engines/v1",
        probe=probe_docker,
        needs_api_key=False,
        notes="Enable Model Runner in Docker Desktop and pull a model with `docker model pull`.",
    ),
    "remote": Provider(
        key="remote",
        label="Remote (OpenAI-compatible)",
        default_base_url="https://api.openai.com/v1",
        chat_base_url_suffix="",
        probe=probe_remote,
        needs_api_key=True,
        notes="Any OpenAI-compatible API. Examples: api.openai.com/v1, api.together.xyz/v1, api.groq.com/openai/v1.",
    ),
}


# ---------------------------------------------------------------------------
# Document parsers
# ---------------------------------------------------------------------------


def _parse_txt(raw: bytes) -> str:
    for enc in ("utf-8", "utf-16", "latin-1"):
        try:
            return raw.decode(enc)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace")


def _parse_pdf(raw: bytes) -> str:
    reader = PdfReader(io.BytesIO(raw))
    pages = []
    for i, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            pages.append(f"--- Page {i} ---\n{text}")
    return "\n\n".join(pages)


def _parse_docx(raw: bytes) -> str:
    doc = DocxDocument(io.BytesIO(raw))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


PARSERS = {
    "txt": _parse_txt,
    "md": _parse_txt,
    "markdown": _parse_txt,
    "pdf": _parse_pdf,
    "docx": _parse_docx,
}


def parse_document(name: str, raw: bytes) -> str:
    ext = Path(name).suffix.lower().lstrip(".")
    parser = PARSERS.get(ext)
    if parser is None:
        raise ValueError(f"Unsupported file type: .{ext}")
    return parser(raw)


# ---------------------------------------------------------------------------
# Streamlit UI
# ---------------------------------------------------------------------------

st.set_page_config(page_title="Document Summarizer", layout="wide")

ss = st.session_state
ss.setdefault("doc_text", None)
ss.setdefault("filename", None)
ss.setdefault("summary", "")
ss.setdefault("models_for_provider", {})
ss.setdefault("context", "")

SAMPLE_DOC = Path(__file__).parent / "sample_document.md"


with st.expander("Settings", expanded=ss.doc_text is None):
    model_col, doc_col = st.columns(2, gap="large")

    with model_col:
        st.markdown("**Model**")
        provider_key = st.selectbox(
            "Provider",
            options=list(PROVIDERS.keys()),
            format_func=lambda k: PROVIDERS[k].label,
            key="provider_key",
        )
        provider = PROVIDERS[provider_key]
        base_url = st.text_input(
            "Base URL", value=provider.default_base_url, key=f"base_{provider_key}"
        )
        api_key = ""
        if provider.needs_api_key:
            api_key = st.text_input(
                "API key", value="", type="password", key=f"key_{provider_key}"
            )
        st.caption(provider.notes)
        if provider.needs_api_key:
            st.warning(
                "⚠ Using a remote provider sends the full document text off "
                "this device to a third-party API."
            )

        if st.button("Scan for running models", use_container_width=True):
            try:
                models = provider.probe(base_url, api_key)
                ss.models_for_provider[provider_key] = models
                if models:
                    st.success(f"Found {len(models)} model(s).")
                else:
                    st.warning("Provider responded but reported no loaded models.")
            except requests.exceptions.ConnectionError:
                st.error(f"Could not reach {base_url}. Is the provider running?")
            except Exception as e:
                st.error(f"Scan failed: {e}")

        available = ss.models_for_provider.get(provider_key, [])
        model = st.selectbox(
            "Model",
            options=available or ["— scan to populate —"],
            disabled=not available,
            key=f"model_{provider_key}",
        )

        st.markdown("**Summary style**")
        style = st.selectbox(
            "Length",
            options=["Short (3-5 bullets)", "Medium (1-2 paragraphs)", "Detailed (sectioned)"],
            index=1,
            key="style",
        )

    with doc_col:
        st.markdown("**Document**")
        upload = st.file_uploader(
            "Upload a document",
            type=["txt", "md", "markdown", "pdf", "docx"],
            accept_multiple_files=False,
        )
        if upload is not None and upload.name != ss.filename:
            try:
                text = parse_document(upload.name, upload.read())
                if not text.strip():
                    st.error("Could not extract any text from this document.")
                else:
                    ss.doc_text = text
                    ss.filename = upload.name
                    ss.summary = ""
                    st.toast(f"Loaded {upload.name} — {len(text):,} characters")
                    st.rerun()
            except Exception as e:
                st.error(f"Could not parse document: {e}")

        if SAMPLE_DOC.exists() and st.button(
            "Load sample document",
            use_container_width=True,
            help=f"Load bundled {SAMPLE_DOC.name} for a quick demo.",
        ):
            ss.doc_text = SAMPLE_DOC.read_text(encoding="utf-8")
            ss.filename = SAMPLE_DOC.name
            ss.summary = ""
            st.toast(f"Loaded {SAMPLE_DOC.name} — {len(ss.doc_text):,} characters")
            st.rerun()


doc_text: str | None = ss.doc_text

if doc_text is None:
    st.info("Upload a document in **Settings** above — or click **Load sample document** to try the bundled example.")
    st.stop()

doc_col, summary_col = st.columns([1, 1], gap="large")

with doc_col:
    st.subheader(ss.filename)
    st.caption(f"{len(doc_text):,} characters")
    st.text_area("Extracted text", value=doc_text, height=480, disabled=True, label_visibility="collapsed")

with summary_col:
    st.subheader("Summary")
    summary_container = st.container(height=360)

    context = st.text_area(
        "Additional context (optional)",
        key="context",
        placeholder="e.g. \"This is a board memo — focus on financial risks.\" or \"Audience is non-technical. Define jargon.\"",
        help="Extra instructions or background that the model should consider when summarizing.",
        height=100,
    )

    button_slot = st.empty()
    generate = button_slot.button(
        "Generate summary", type="primary", use_container_width=True, key="gen_btn"
    )

    if generate:
        if not available:
            st.warning("Pick a provider and scan for a model first.")
            st.stop()

        button_slot.button(
            "Generating…",
            type="primary",
            use_container_width=True,
            disabled=True,
            key="gen_btn_disabled",
        )

        style_directive = {
            "Short (3-5 bullets)": "Produce a tight summary as 3-5 bullet points covering the most important takeaways.",
            "Medium (1-2 paragraphs)": "Produce a 1-2 paragraph summary that captures the document's purpose and key points.",
            "Detailed (sectioned)": "Produce a detailed, sectioned summary with headings for each major topic and bullet points beneath each heading.",
        }[style]

        system = (
            "You are an expert summarizer. Read the provided document carefully and produce a faithful, "
            "non-fabricated summary. Preserve technical accuracy, named entities, and numbers. "
            f"{style_directive}"
        )
        context_block = (
            f"Additional context from the user (treat as instructions, not as part of the document):\n{context.strip()}\n\n"
            if context.strip()
            else ""
        )
        user = (
            f"{context_block}"
            f"Document: {ss.filename}\n\n---\n{doc_text}\n---\n\n"
            "Summarize the document above."
        )

        client = OpenAI(
            base_url=f"{base_url.rstrip('/')}{provider.chat_base_url_suffix}",
            api_key=api_key if (provider.needs_api_key and api_key) else "not-needed",
        )

        with summary_container:
            placeholder = st.empty()
            collected = []
            with st.spinner("Summarizing…"):
                try:
                    stream = client.chat.completions.create(
                        model=model,
                        messages=[
                            {"role": "system", "content": system},
                            {"role": "user", "content": user},
                        ],
                        stream=True,
                    )
                    for chunk in stream:
                        delta = chunk.choices[0].delta.content if chunk.choices else None
                        if delta:
                            collected.append(delta)
                            placeholder.markdown("".join(collected))
                    ss.summary = "".join(collected) or "(empty response)"
                except Exception as e:
                    ss.summary = f"**Error talking to {provider.label}:** {e}"
                    placeholder.markdown(ss.summary)
        st.rerun()
    elif ss.summary:
        with summary_container:
            st.markdown(ss.summary)
    else:
        with summary_container:
            st.caption("Click **Generate summary** to summarize this document.")
